# app/services/fj_httpx_fallback.py
"""
Fallback para consultas de Función Judicial usando API directa (HTTPX).

VERSIÓN V3 - CORRECCIÓN CRÍTICA:
- Distingue entre "sin resultados" (API 200 + lista vacía) vs "error API" (500, timeout, etc.)
- NO genera reporte cuando hay error de API
- Retorna scenario='api_error' para que daemon resetee cliente a Pendiente

Características:
- Consulta hasta 20 páginas
- Genera DOCX con tablas formateadas
- Convierte fechas UTC → Ecuador (UTC-5)
- Guarda en sri_ruc_output/reports/
"""

import httpx
from docx import Document
from datetime import datetime, timedelta, timezone
import os
from typing import Optional, List, Dict, Any, Tuple
import traceback
from app.services.word_utils import agregar_linea_clave_valor, agregar_titulo_principal, configurar_documento

# ===== CONFIGURACIÓN =====
API_BASE_URL = "https://api.funcionjudicial.gob.ec/EXPEL-CONSULTA-CAUSAS-SERVICE"
PAGE_SIZE = 10
MAX_PAGES = 20
REPORTS_DIR = "sri_ruc_output/reports"

# Crear directorio si no existe
os.makedirs(REPORTS_DIR, exist_ok=True)


def log(msg: str):
    """Logging con timestamp para HTTPX fallback"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[HTTPX FALLBACK {timestamp}] {msg}")


def _convertir_fecha_utc_a_ecuador(fecha_str: str) -> str:
    """
    Convierte fecha UTC a hora de Ecuador (UTC-5).
    
    Args:
        fecha_str: Fecha en formato ISO (ej: "2025-11-17T00:00:00")
        
    Returns:
        Fecha formateada como dd/mm/yyyy
    """
    try:
        if not fecha_str:
            return "N/A"
        
        # Parsear fecha
        if "T" in fecha_str:
            dt_utc = datetime.fromisoformat(fecha_str.replace("Z", "+00:00"))
        else:
            dt_utc = datetime.fromisoformat(fecha_str)
        
        # Asegurar que tenga timezone UTC
        if dt_utc.tzinfo is None:
            dt_utc = dt_utc.replace(tzinfo=timezone.utc)
        
        # Convertir a Ecuador (UTC-5)
        dt_ec = dt_utc - timedelta(hours=5)
        
        # Retornar formateado
        return dt_ec.strftime("%d/%m/%Y")
    except Exception as e:
        log(f"⚠️ Error convirtiendo fecha '{fecha_str}': {e}")
        return fecha_str[:10] if len(fecha_str) >= 10 else "N/A"


def _consultar_pagina_api(nombre_buscado: str, page: int) -> Tuple[Optional[List[Dict[str, Any]]], bool]:
    """
    Consulta una página de la API de Función Judicial.
    
    Args:
        nombre_buscado: Nombre del demandado a buscar
        page: Número de página (1-based)
        
    Returns:
        Tupla (resultados, fue_exitoso):
        - (lista_datos, True) → API respondió 200, hay datos
        - ([], True) → API respondió 200, sin datos (lista vacía)
        - (None, False) → Error de API (500, timeout, etc.)
    """
    try:
        url = (
            f"{API_BASE_URL}/"
            f"api/consulta-causas/informacion/buscarCausas"
            f"?page={page}&size={PAGE_SIZE}"
        )
        
        payload = {
            "numeroCausa": "",
            "actor": {
                "cedulaActor": "",
                "nombreActor": ""
            },
            "demandado": {
                "cedulaDemandado": "",
                "nombreDemandado": nombre_buscado
            },
            "provincia": "",
            "numeroFiscalia": "",
            "recaptcha": "",
            "first": page,
            "pageSize": PAGE_SIZE
        }
        
        # ===== DEBUG: Ver exactamente qué se envía =====
        log(f"🔍 DEBUG - Payload a enviar:")
        log(f"   URL: {url}")
        log(f"   nombreDemandado: '{nombre_buscado}'")
        log(f"   Repr: {repr(nombre_buscado)}")
        log(f"   Longitud: {len(nombre_buscado)} caracteres")
        
        headers = {
            "Content-Type": "application/json",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        }
        
        # Usar httpx con timeout
        with httpx.Client(timeout=30.0) as client:
            response = client.post(url, json=payload, headers=headers)
        
        # ===== DEBUG: Ver respuesta =====
        log(f"🔍 DEBUG - Respuesta API:")
        log(f"   Status: {response.status_code}")
        
        # ✅ CORRECCIÓN: Distinguir entre error de API vs sin resultados
        if response.status_code != 200:
            log(f"⚠️ API retornó status {response.status_code}")
            # Intentar ver el cuerpo del error
            try:
                error_body = response.text[:500]  # Primeros 500 chars
                log(f"   Body error: {error_body}")
            except:
                pass
            return (None, False)  # Error de API
        
        # Parsear respuesta
        data = response.json()
        
        # ===== DEBUG: Ver datos recibidos =====
        log(f"🔍 DEBUG - Datos recibidos:")
        log(f"   Tipo: {type(data).__name__}")
        if isinstance(data, dict):
            log(f"   Keys: {list(data.keys())}")
            resultados = data.get("data", [])
        elif isinstance(data, list):
            resultados = data
        else:
            log(f"   Formato inesperado")
            return (None, False)  # Respuesta inválida
        
        log(f"   Cantidad resultados: {len(resultados) if resultados else 0}")
        
        # ✅ API respondió 200 exitosamente
        if resultados:
            return (resultados, True)  # Datos encontrados
        else:
            return ([], True)  # Sin datos, pero API funcionó correctamente
        
    except httpx.TimeoutException:
        log(f"⚠️ Timeout consultando página {page}")
        return (None, False)  # Error de red
    except Exception as e:
        log(f"⚠️ Error consultando API página {page}: {e}")
        return (None, False)  # Error general


# Función auxiliar para formatear nombre completo de encabezado
def _formatear_nombre_completo(apellidos: str, nombres: str) -> str:
    """Combina APELLIDOS + NOMBRES, retorna 'NO APLICA' si ambos vacíos"""
    apellidos = (apellidos or "").strip()
    nombres = (nombres or "").strip()
    nombre_completo = f"{apellidos} {nombres}".strip()
    return nombre_completo if nombre_completo else "NO APLICA"


def _valor_o_no_aplica(valor: str) -> str:
    """Retorna el valor si existe, sino 'NO APLICA'"""
    if valor and str(valor).strip() and str(valor).strip().upper() not in ['N/A', 'NA', 'NONE', '']:
        return str(valor).strip()
    return "NO APLICA"


def generar_reporte_httpx(
    nombre_cliente: str,
    job_id: str,
    meta: Dict[str, Any] = None
) -> Tuple[Optional[str], Dict[str, Any]]:
    """
    Genera reporte DOCX consultando API de Función Judicial directamente.
    
    ✅ VERSIÓN V3 - CORRECCIÓN CRÍTICA:
    - Distingue entre "sin resultados" vs "error API"
    - NO genera reporte cuando hay error de API (500, timeout, etc.)
    - Retorna scenario='api_error' para que daemon resetee cliente a Pendiente
    
    Args:
        nombre_cliente: Nombre completo del cliente (ej: "PAMELA ALEXANDRA CASTRO DEL POZO")
        job_id: ID único del proceso (ej: "daemon_8d8c1f044264")
        meta: Diccionario con datos del cliente para el encabezado profesional
        
    Returns:
        Tupla (ruta_reporte, resultado_dict):
        - (ruta, {scenario: 'results_found'}) → Reporte con datos
        - (ruta, {scenario: 'no_results'}) → Reporte sin datos (API OK pero sin procesos)
        - (None, {scenario: 'api_error'}) → Error de API, NO se genera reporte
    """
    try:
        log(f"🌐 Iniciando consulta API para: {nombre_cliente}")
        
        # ===== FASE 1: CONSULTAR API =====
        # Primero consultamos la página 1 para ver si la API responde
        
        log(f"📄 Consultando página 1...")
        resultados_p1, api_exitosa = _consultar_pagina_api(nombre_cliente, 1)
        
        # ✅ CORRECCIÓN CRÍTICA: Si la API falló, NO generar reporte
        if not api_exitosa:
            log(f"❌ Error de API en página 1 - NO se generará reporte")
            return (None, {
                "scenario": "api_error",
                "total_procesos": 0,
                "total_paginas": 0,
                "mensaje": "Error de API (500, timeout, etc.) - Cliente debe reintentarse"
            })
        
        # ===== FASE 2: API RESPONDIÓ CORRECTAMENTE =====
        # Ahora sí creamos el documento
        
        doc = Document()
        configurar_documento(doc)  # ✅ Aplicar estilos globales
        agregar_titulo_principal(doc, "Revisión de Función Judicial")
        doc.add_paragraph("")  # Espacio después del título
        
        # ===== ENCABEZADO PROFESIONAL (7 CAMPOS CLAVE-VALOR) =====
        if meta:
            # Fecha de consulta
            agregar_linea_clave_valor(doc, "FECHA DE CONSULTA", datetime.now().strftime("%d/%m/%Y"))
            
            # Titular
            agregar_linea_clave_valor(doc, "NOMBRE Y APELLIDO DEL TITULAR", _valor_o_no_aplica(meta.get('cliente_nombre')))
            agregar_linea_clave_valor(doc, "NUMERO DE CEDULA DEL TITULAR", _valor_o_no_aplica(meta.get('cliente_cedula')))
            
            # Cónyuge: combinar APELLIDOS + NOMBRES
            nombre_conyuge_completo = _formatear_nombre_completo(
                meta.get('apellidos_conyuge', ''),
                meta.get('nombres_conyuge', '')
            )
            agregar_linea_clave_valor(doc, "NOMBRE DEL CONYUGE", nombre_conyuge_completo)
            agregar_linea_clave_valor(doc, "CEDULA DEL CONYUGE", _valor_o_no_aplica(meta.get('cedula_conyuge')))
            
            # Codeudor: combinar APELLIDOS + NOMBRES
            nombre_codeudor_completo = _formatear_nombre_completo(
                meta.get('apellidos_codeudor', ''),
                meta.get('nombres_codeudor', '')
            )
            agregar_linea_clave_valor(doc, "NOMBRE DE CODEUDOR", nombre_codeudor_completo)
            agregar_linea_clave_valor(doc, "CEDULA DEL CODEUDOR", _valor_o_no_aplica(meta.get('cedula_codeudor')))
        else:
            # Fallback: si no hay meta, usar datos mínimos
            agregar_linea_clave_valor(doc, "FECHA DE CONSULTA", datetime.now().strftime("%d/%m/%Y"))
            agregar_linea_clave_valor(doc, "NOMBRE Y APELLIDO DEL TITULAR", nombre_cliente)
            agregar_linea_clave_valor(doc, "NUMERO DE CEDULA DEL TITULAR", "NO APLICA")
            agregar_linea_clave_valor(doc, "NOMBRE DEL CONYUGE", "NO APLICA")
            agregar_linea_clave_valor(doc, "CEDULA DEL CONYUGE", "NO APLICA")
            agregar_linea_clave_valor(doc, "NOMBRE DE CODEUDOR", "NO APLICA")
            agregar_linea_clave_valor(doc, "CEDULA DEL CODEUDOR", "NO APLICA")
        
        doc.add_paragraph("")  # Espacio después del encabezado
        
        # ===== FASE 3: PROCESAR RESULTADOS =====
        
        contador = 1
        total_resultados = 0
        alguna_pagina_con_datos = False
        pagina_actual = 1
        
        # Procesar página 1 (ya la consultamos)
        if resultados_p1:
            alguna_pagina_con_datos = True
            total_resultados += len(resultados_p1)
            
            # Agregar encabezado de página
            doc.add_heading(f"Página 1", level=2)
            
            # Crear tabla con 5 columnas
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            
            # Encabezados
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "No."
            hdr_cells[1].text = "Fecha de ingreso"
            hdr_cells[2].text = "No. proceso"
            hdr_cells[3].text = "Acción / Infracción"
            hdr_cells[4].text = "Movimientos del Proceso"
            
            # Llenar tabla con datos
            for caso in resultados_p1:
                row = table.add_row().cells
                
                # Convertir fecha
                fecha_api = caso.get("fechaIngreso", "")
                fecha_formato = _convertir_fecha_utc_a_ecuador(fecha_api) if fecha_api else "N/A"
                
                # Llenar celdas
                row[0].text = str(contador)
                row[1].text = fecha_formato
                row[2].text = str(caso.get("idJuicio", "N/A"))
                row[3].text = str(caso.get("nombreDelito", "N/A"))
                row[4].text = "🗂️"
                
                contador += 1
            
            doc.add_paragraph("")
        
        # Consultar páginas adicionales (2 en adelante) solo si página 1 tuvo datos
        if alguna_pagina_con_datos:
            for page in range(2, MAX_PAGES + 1):
                log(f"📄 Consultando página {page}...")
                
                resultados, api_ok = _consultar_pagina_api(nombre_cliente, page)
                
                # Si la API falló en páginas posteriores, detenemos pero NO invalidamos
                # porque ya tenemos datos de páginas anteriores
                if not api_ok:
                    log(f"⚠️ Error en página {page}, continuando con datos obtenidos...")
                    break
                
                if not resultados:
                    # Lista vacía = no hay más páginas
                    log(f"📭 Página {page} sin resultados, finalizando")
                    break
                
                # Sí hay datos
                pagina_actual = page
                total_resultados += len(resultados)
                
                # Agregar encabezado de página
                doc.add_heading(f"Página {page}", level=2)
                
                # Crear tabla con 5 columnas
                table = doc.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                
                # Encabezados
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "No."
                hdr_cells[1].text = "Fecha de ingreso"
                hdr_cells[2].text = "No. proceso"
                hdr_cells[3].text = "Acción / Infracción"
                hdr_cells[4].text = "Movimientos del Proceso"
                
                # Llenar tabla con datos
                for caso in resultados:
                    row = table.add_row().cells
                    
                    fecha_api = caso.get("fechaIngreso", "")
                    fecha_formato = _convertir_fecha_utc_a_ecuador(fecha_api) if fecha_api else "N/A"
                    
                    row[0].text = str(contador)
                    row[1].text = fecha_formato
                    row[2].text = str(caso.get("idJuicio", "N/A"))
                    row[3].text = str(caso.get("nombreDelito", "N/A"))
                    row[4].text = "Movimientos del Proceso"
                    
                    contador += 1
                
                doc.add_paragraph("")
        
        # ===== FASE 4: DETERMINAR ESCENARIO Y GUARDAR =====
        
        if alguna_pagina_con_datos:
            scenario = "results_found"
            mensaje = f"Se encontraron {total_resultados} procesos judiciales en {pagina_actual} página(s)"
        else:
            scenario = "no_results"
            mensaje = "NO SE ENCONTRARON PROCESOS JUDICIALES"
            # Agregar mensaje al documento
            doc.add_paragraph(mensaje)
        
        # Guardar documento
        nombre_archivo = f"reporte_{nombre_cliente.replace(' ', '_')}.docx"
        ruta_completa = os.path.join(REPORTS_DIR, nombre_archivo)
        
        try:
            doc.save(ruta_completa)
            log(f"✅ Reporte DOCX generado: {ruta_completa}")
            log(f"   - Escenario: {scenario}")
            log(f"   - Total procesos: {total_resultados}")
            log(f"   - Páginas: {pagina_actual}")
        except Exception as e:
            log(f"❌ Error guardando documento: {e}")
            return (None, {
                "scenario": "error",
                "total_procesos": 0,
                "total_paginas": 0,
                "mensaje": f"Error guardando documento: {str(e)}"
            })
        
        # ✅ Retornar resultado exitoso
        resultado = {
            "scenario": scenario,
            "total_procesos": total_resultados,
            "total_paginas": pagina_actual,
            "mensaje": mensaje
        }
        
        return (ruta_completa, resultado)
        
    except Exception as e:
        log(f"❌ Error generando reporte HTTPX: {e}")
        traceback.print_exc()
        
        # Error crítico
        return (None, {
            "scenario": "error",
            "total_procesos": 0,
            "total_paginas": 0,
            "mensaje": f"Error crítico: {str(e)}"
        })