from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ===== CONFIGURACIÓN GLOBAL =====
FUENTE_PRINCIPAL = "Times New Roman"
TAMAÑO_NORMAL = Pt(12)
TAMAÑO_TITULO = Pt(16)


def configurar_documento(doc):
    """
    Configura estilos globales del documento:
    - Fuente: Times New Roman
    - Tamaño: 12pt
    - Márgenes: 2.54cm (estándar)
    """
    # Configurar márgenes
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    
    # Configurar estilo Normal (afecta a todo el documento)
    style = doc.styles["Normal"]
    style.font.name = FUENTE_PRINCIPAL
    style.font.size = TAMAÑO_NORMAL
    
    # Configurar estilos de Heading también
    for i in range(1, 4):  # Heading 1, 2, 3
        try:
            heading_style = doc.styles[f"Heading {i}"]
            heading_style.font.name = FUENTE_PRINCIPAL
        except:
            pass
    
    return doc


def agregar_linea_clave_valor(doc, clave, valor):
    """
    Agrega una línea con formato CLAVE: valor
    - Clave en MAYÚSCULAS y negrilla
    - Valor sin negrilla
    """
    p = doc.add_paragraph()
    
    # Añadir la clave en mayúsculas y negrilla
    run_clave = p.add_run(clave.upper() + ": ")
    run_clave.bold = True
    run_clave.font.name = FUENTE_PRINCIPAL
    
    # Añadir el valor sin negrilla
    run_valor = p.add_run(valor)
    run_valor.font.name = FUENTE_PRINCIPAL
    
    return p


def agregar_titulo_principal(doc, texto):
    """
    Agrega un título principal centrado, en mayúsculas, negro y Times New Roman.
    """
    titulo = doc.add_heading(texto.upper(), level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Aplicar formato a cada run
    for run in titulo.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        run.font.name = FUENTE_PRINCIPAL
        run.font.size = TAMAÑO_TITULO
    
    return titulo


def agregar_subtitulo(doc, texto):
    """
    Agrega un subtítulo alineado a la izquierda.
    """
    subtitulo = doc.add_heading(texto, level=2)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Aplicar formato
    for run in subtitulo.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        run.font.name = FUENTE_PRINCIPAL
    
    return subtitulo